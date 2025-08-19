from flask import Flask, request, render_template, send_file, jsonify
import os
import tempfile
from werkzeug.utils import secure_filename
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Inches
from PIL import Image
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024

UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in {'pdf', 'docx'}

def pdf_to_docx_simple(pdf_path, output_path):
    """간단한 PDF to DOCX 변환"""
    try:
        images = convert_from_path(pdf_path, dpi=150)
        doc = Document()
        
        for i, image in enumerate(images):
            with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as temp_file:
                image.save(temp_file.name, 'JPEG', quality=85)
                doc.add_picture(temp_file.name, width=Inches(6))
                
                if i < len(images) - 1:
                    doc.add_page_break()
                    
                os.unlink(temp_file.name)
        
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"PDF to DOCX 오류: {e}")
        return False

def docx_to_pdf_simple(docx_path, output_path):
    """간단한 DOCX to PDF 변환"""
    try:
        doc = Document(docx_path)
        c = canvas.Canvas(output_path, pagesize=A4)
        width, height = A4
        y_position = height - 50
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                c.drawString(50, y_position, paragraph.text[:100])
                y_position -= 20
                if y_position < 50:
                    c.showPage()
                    y_position = height - 50
        
        c.save()
        return True
    except Exception as e:
        print(f"DOCX to PDF 오류: {e}")
        return False

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/convert', methods=['POST'])
def convert_file():
    try:
        print("변환 요청 받음")
        
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': '파일이 없습니다.'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'error': '파일명이 없습니다.'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'success': False, 'error': 'PDF 또는 DOCX 파일만 가능합니다.'}), 400
        
        filename = secure_filename(file.filename)
        file_path = os.path.join(UPLOAD_FOLDER, filename)
        file.save(file_path)
        
        file_ext = filename.rsplit('.', 1)[1].lower()
        
        if file_ext == 'pdf':
            output_filename = filename.rsplit('.', 1)[0] + '.docx'
            output_path = os.path.join(OUTPUT_FOLDER, output_filename)
            success = pdf_to_docx_simple(file_path, output_path)
        else:  # docx
            output_filename = filename.rsplit('.', 1)[0] + '.pdf'
            output_path = os.path.join(OUTPUT_FOLDER, output_filename)
            success = docx_to_pdf_simple(file_path, output_path)
        
        # 임시 파일 삭제
        if os.path.exists(file_path):
            os.remove(file_path)
        
        if success and os.path.exists(output_path):
            return send_file(output_path, as_attachment=True, download_name=output_filename)
        else:
            return jsonify({'success': False, 'error': '변환 실패'}), 500
            
    except Exception as e:
        print(f"변환 오류: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

if __name__ == '__main__':
    print("🚀 간단한 PDF ↔ DOCX 변환기 시작")
    app.run(debug=True, host='0.0.0.0', port=5000)